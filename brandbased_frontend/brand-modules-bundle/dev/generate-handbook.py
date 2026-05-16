"""
BrandBased — generate the product handbook as a .docx.

Run from the repo root:
    python3 BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-modules-bundle/dev/generate-handbook.py

The script depends on `python-docx`, installed locally into
`BB-Entire-Frontend-Modals-and-Brand-Dashbaord/.docgen-venv` (see the
session that produced this file). Output lands at:

    BB-Entire-Frontend-Modals-and-Brand-Dashbaord/BrandBased-Product-Handbook.docx

Content is hard-coded prose summarising every dashboard page, every
Freemium-vs-Premium fork, the shared chrome / popups, and the dev tools.
"""

from __future__ import annotations

import os
import sys
from pathlib import Path


HERE = Path(__file__).resolve().parent
BUNDLE_ROOT = HERE.parent                  # .../brand-modules-bundle
BB_ROOT = BUNDLE_ROOT.parent               # .../BB-Entire-Frontend-Modals-and-Brand-Dashbaord
VENDOR = BB_ROOT / ".docgen-venv"
if VENDOR.exists():
    sys.path.insert(0, str(VENDOR))

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_PURPLE = RGBColor(0x10, 0x2F, 0xF5)
INK = RGBColor(0x12, 0x14, 0x1C)
MUTED = RGBColor(0x55, 0x5B, 0x68)
ACCENT_AMBER = RGBColor(0xD6, 0x89, 0x10)


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_styled_run(paragraph, text, *, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_heading(doc, text, level=1, color=BRAND_PURPLE, after_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    sizes = {0: 28, 1: 20, 2: 15, 3: 12}
    add_styled_run(p, text, bold=True, color=color, size=sizes.get(level, 12))
    return p


def add_body(doc, text, *, bold=False, italic=False, color=INK, size=11, after_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    add_styled_run(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


def add_kv_paragraph(doc, label, value):
    """Bold label, regular value, on one line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_styled_run(p, f"{label}: ", bold=True, color=INK, size=11)
    add_styled_run(p, value, color=INK, size=11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    add_styled_run(p, text, size=11)
    return p


def add_two_col_table(doc, headers, rows, *, header_fill="102FF5", header_text=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    table.allow_autofit = False
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
        hdr[i].paragraphs[0].text = ""
        add_styled_run(hdr[i].paragraphs[0], h, bold=True, color=header_text, size=10)
    for r_i, row in enumerate(rows, start=1):
        cells = table.rows[r_i].cells
        for c_i, val in enumerate(row):
            cells[c_i].paragraphs[0].text = ""
            add_styled_run(cells[c_i].paragraphs[0], str(val), size=10)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    add_styled_run(p, "—" * 30, color=MUTED, size=9)


# =====================================================================
# Document content
# =====================================================================

def build():
    doc = Document()

    # ---- Document defaults ------------------------------------------
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # ---- Cover -------------------------------------------------------
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_run(cover_title, "BrandBased — Brand Console", bold=True, color=BRAND_PURPLE, size=30)

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_run(cover_sub, "Product Handbook — page-by-page reference, Freemium & Premium flows.", color=MUTED, size=13)

    cover_meta = doc.add_paragraph()
    add_styled_run(cover_meta, "Scope: ", bold=True, size=10, color=MUTED)
    add_styled_run(cover_meta, "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/", italic=True, size=10, color=MUTED)
    add_styled_run(cover_meta, "  ·  Audience: product, design, backend, support.", size=10, color=MUTED)

    doc.add_paragraph()

    # ---- Section: Plan tiers ----------------------------------------
    add_heading(doc, "1. Plan tiers at a glance", level=1)
    add_body(doc,
        "BrandBased has two account tiers. Both share the same dashboard chrome (header, "
        "sidebar, footer, account menu). The difference is what each tier can do inside the "
        "modules and which iframe pages they're routed to."
    )

    add_heading(doc, "Premium", level=2)
    add_bullet(doc, "Up to 12 Premium brands and 12 Freemium brands per account (24 total).")
    add_bullet(doc, "Full Theme Styles & Theme Settings (7 slots), Products (7 slots), Hotspots, AI Logic, Ad Campaigns, Metrics, Payments, Scheduling, Platforms.")
    add_bullet(doc, "Standalone Premium popup modal renders the customer's last-saved brand + theme + products.")
    add_bullet(doc, "Account chip in the header reads \"Account & Subscriptions\".")
    add_bullet(doc, "\"Get Premium\" CTA in the header is hidden for Premium users.")

    add_heading(doc, "Freemium", level=2)
    add_bullet(doc, "Up to 12 Freemium brands per account. No Premium brand quota.")
    add_bullet(doc, "Restricted module set — only Brands, Brand Settings, Theme Styles, Hotspots, and the upload/verification sub-flows.")
    add_bullet(doc, "Standalone Freemium popup modal renders the customer's last-saved brand + the active theme slot only (products ignored).")
    add_bullet(doc, "Account chip in the header reads just \"Account\" (the \"& Subscriptions\" suffix is hidden).")
    add_bullet(doc, "\"Get Premium\" CTA appears in the header and opens the Unlock Premium upgrade gate.")
    add_bullet(doc, "Premium-only routes (Products, AI Logic, Ad Campaigns, etc.) are covered by a full-viewport Unlock Premium gate instead of loading the module.")

    add_body(doc,
        "While the production plan-tier flag is being wired, both tiers are simulated locally via "
        "the dev tool's \"Brand Console — dashboard tier preview\" control. It writes "
        "localStorage[\"bbDevDashSimulateFreemium\"] = \"1\" (Freemium) or removes the key (Premium). "
        "The dashboard's gate script reflects this state by toggling body.bb-dev-freemium-mode on "
        "the dashboard document, which CSS uses to swap chrome (Get Premium button, Account chip).",
        italic=True, color=MUTED, size=10,
    )
    add_divider(doc)

    # ---- Section: Dashboard chrome ----------------------------------
    add_heading(doc, "2. Dashboard chrome", level=1)
    add_body(doc,
        "Every module page loads inside the dashboard's iframe (#bbDashFrame inside .content-block). "
        "The chrome around it stays mounted across navigations:"
    )

    add_heading(doc, "Header", level=2)
    add_bullet(doc, "Animated blue accent line under the header bar.")
    add_bullet(doc, "Account chip (top-right). Click opens the Account dropdown.")
    add_bullet(doc, "Get Premium pill (Freemium only).")
    add_bullet(doc, "Heading + decorative video (Start Now only on desktop).")

    add_heading(doc, "Account dropdown", level=2)
    add_body(doc, "Items, in order:")
    add_bullet(doc, "Subscriptions — closes the dropdown; on Freemium also surfaces the Unlock Premium gate.")
    add_bullet(doc, "Preferences — opens the Preferences modal (see §4).")
    add_bullet(doc, "Theme Light / Theme Dark toggle.")
    add_bullet(doc, "Logout button.")

    add_heading(doc, "Sidebar (desktop) / Footer slide-up (mobile)", level=2)
    add_body(doc, "Nav items wired via data-route. Each item maps to a route in dashboard-router.js. The router:")
    add_bullet(doc, "Swaps the iframe src to the route's URL.")
    add_bullet(doc, "Shows the loader (blurred glass + breathing B + sweep line) for at least MIN_LOADER_MS.")
    add_bullet(doc, "Updates ?page=<id> in the address bar via history.replaceState.")
    add_bullet(doc, "Strips per-module decorative backgrounds (.light-shard-ui-base, .bb-bg-shards, .video-bg-content-header) so the dashboard's own background shows through cleanly.")
    add_bullet(doc, "Hard-refreshes the dashboard at ?page=<id> for Theme Styles + Products (their inner-scroll layout doesn't recover from in-iframe nav).")

    add_heading(doc, "Footer", level=2)
    add_bullet(doc, "B icon + © BrandBased pinned to the left edge (5px in).")
    add_bullet(doc, "Footer is right-anchored; the B + copyright are nudged left independently.")
    add_divider(doc)

    # ---- Section: Pages ---------------------------------------------
    add_heading(doc, "3. Pages — how each module works", level=1)
    add_body(doc,
        "All page URLs below are relative to brand-modules-bundle/. Routes are declared in "
        "brand-console-final/js/dashboard-router.js (the ROUTES + ROUTE_ALIASES tables)."
    )

    pages = [
        # (heading, sub-rows)
        (
            "Start Now",
            "start-now/Start-Now.html",
            [
                "Entry point for new accounts. Two CTAs: Go Premium and Start FREE.",
                "Go Premium → routes to the Premium logo upload page (logo-upload/Logo-upload-and-Crop-module.html).",
                "Start FREE → routes to the Freemium logo upload page (freemium/Freemium-Logo-upload-and-Crop-module.html), which is a duplicated file pinned to the Freemium flow.",
                "The dashboard's #content-heading + .video-bg-content-header are only shown on this page (on desktop).",
                "Excluded from the Unlock Premium gate — Freemium users still see it.",
            ],
        ),
        (
            "Brands",
            "brands/Brands.html",
            [
                "Grid of brand circles, up to 8 across on desktop.",
                "Each brand card has: a circular brand mark, a name, a Publish / Unpublish action, and a Delete (×) action.",
                "Premium pill (top-left) appears when the dev \"View whether this brand is using a Free or Premium plan\" toggle is set — uses premium-start.svg with a purple highlight and a white star.",
                "Unverified badge (bb-brands-badge--unverified) appears on every brand when the dev \"Brand Verification\" toggle is set to Fail.",
                "Add Brand circle → opens a small glass popover with two options: Freemium and Premium.",
                "Add → Freemium navigates to ../freemium/Freemium-Logo-upload-and-Crop-module.html (matches Start FREE).",
                "Add → Premium navigates to ../logo-upload/Logo-upload-and-Crop-module.html, BUT when dev Freemium mode is on, it triggers the Unlock Premium gate instead.",
                "Delete (×) → opens the full-viewport Delete confirm popup (red B icon, red breathing halo, full-page blur). Confirming fires the demo's delete action.",
                "Excluded from the Unlock Premium gate — Freemium users still need access to their own brand entries.",
            ],
        ),
        (
            "Brand Settings",
            "Brand-Settings-Module.html  /  freemium/Freemium-Brand-Settings.html",
            [
                "Premium and Freemium have separate HTML files. Both render brand details and link to Theme Styles / Hotspots / Products in the purple side panel.",
                "Freemium version's purple panel omits Premium-only links and routes the Theme Styles link to ./Freemium-Theme-Styles.html and Hotspots to ./Freemium-Hotspots.html.",
                "Continue to Brand Verification on the upload/verify sub-flow ends here on Premium and at Freemium-Brand-Settings.html on Freemium.",
            ],
        ),
        (
            "Theme Styles & Theme Settings",
            "Brand-Theme-Settings-Module.html  /  freemium/Freemium-Theme-Styles.html",
            [
                "Premium has 7 theme slots (themeSlot:0 … themeSlot:6). Per slot the user can set: accent colour, background (solid / gradient / image / video), uploaded background image/video, buy-now CTA label, typography, ads/theme preview state, admin gallery, recent colour swatches.",
                "Active slot is stored in bbTheme:activeSlot:v1 and is the slot the popup renders.",
                "Freemium-Theme-Styles.html is a separate page with 1 slot per brand. UI controls are the same; only the slot count differs.",
                "Theme Styles uses a desktop \"fixed page, inner-panel scroll\" layout. The router does a hard refresh of the dashboard at ?page=theme-design when navigating in, so the page lands at the top.",
                "Freemium variant has 20px padding-top to give the iframe a bit of breathing room.",
                "Box shadows on .bb-bts-card are removed across both Premium and Freemium copies of Theme Styles and Products for a flatter, more legible card stack.",
            ],
        ),
        (
            "Products",
            "products/Products.html",
            [
                "Premium-only. 7 slots, mirroring Theme Styles' slot index.",
                "Per slot: array of products (name, price, image, hotspots, etc.) stored under bbProducts:slot:<n>:v1.",
                "Same hard-refresh + inner-scroll behaviour as Theme Styles.",
                "Hidden behind the Unlock Premium gate when the dev Freemium-mode flag is set.",
            ],
        ),
        (
            "Hotspots",
            "hotspots/Hotspots.html  /  freemium/Freemium-Hotspots.html",
            [
                "Premium Hotspots manages interactive product hotspots on uploaded brand assets / themes.",
                "Freemium-Hotspots.html is a separate, simplified file with a single brand asset.",
                "When the user reaches Freemium Hotspots via Freemium Theme Styles' purple \"Enable Hotspots\" CTA, the dashboard router force-reloads the iframe once so the page paints from the top (otherwise it lands partway down).",
            ],
        ),
        (
            "Upload Brand Assets",
            "logo-upload/Logo-upload-and-Crop-module.html  /  freemium/Freemium-Logo-upload-and-Crop-module.html",
            [
                "Light/dark logo uploads with a cropper modal.",
                "Premium copy continues to Premium Meta-Verification.html.",
                "Freemium copy continues to Freemium-Meta-Verification.html — duplicated file means the flow can never \"escape\" back to the Premium page.",
                "Continue button has been rebranded \"Sync\" with the shared sync popup (\"Syncing brand data…\" → \"Synced\").",
                "Uses .bb-lu-tile / .bb-lu-tile--light tiles with a soft 8% drop shadow on the upload buttons; tile caption sits 10px below.",
            ],
        ),
        (
            "Brand Verification",
            "logo-upload/Meta-Verification.html  /  freemium/Freemium-Meta-Verification.html",
            [
                "Side-by-side Light Mode + Dark Mode preview of the verified brand mark.",
                "Default (\"Pass Verify\"): green badges + tick, \"Brand Successfully Verified\" pill, title reads \"Verified Brand Identity\".",
                "Dev toggle \"Brand Verification preview state\" set to Not Verified: red badges + cross, \"Brand Not Verified\" pill, title swaps to \"Unable to Verify Brand. Please contact support@brandbased.ai\" in red.",
                "Verify button uses the shared sync popup with custom labels (\"Verifying\" → \"Verified\").",
                "Continue to Brand Settings → goes to Brand-Settings-Module.html on Premium and freemium/Freemium-Brand-Settings.html on Freemium.",
                "Desktop only: #bbMetaVerificationMount has margin-top: -20px so the verify card sits closer to the page top.",
            ],
        ),
        (
            "AI Logic / Admin-Team-Console-Brand-Settings-AI-Rules",
            "ai-logic/AI-Logic.html  (aliased to Admin-Team-Console-Brand-Settings-AI-Rules.html)",
            [
                "Premium-only. AI rules configuration surface.",
                "Hidden behind the Unlock Premium gate for Freemium users.",
            ],
        ),
        (
            "Ad Campaigns",
            "ad-campaigns/Ad-Campaigns.html",
            [
                "Premium-only campaign manager.",
                "Hidden behind the Unlock Premium gate for Freemium users.",
            ],
        ),
        (
            "Metrics",
            "metrics/Metrics.html",
            [
                "Premium-only analytics dashboard.",
                "Hidden behind the Unlock Premium gate for Freemium users.",
            ],
        ),
        (
            "Payments",
            "payments/Payments.html",
            [
                "Premium-only payments + billing surface.",
                "Hidden behind the Unlock Premium gate for Freemium users.",
            ],
        ),
        (
            "Scheduling",
            "Product-Scheduling-Module.html",
            [
                "Premium-only product scheduling.",
                "Hidden behind the Unlock Premium gate for Freemium users.",
            ],
        ),
        (
            "Platforms",
            "platforms/Platforms.html",
            [
                "Premium-only outbound platform connections.",
                "Hidden behind the Unlock Premium gate for Freemium users.",
            ],
        ),
    ]

    for title, path, bullets in pages:
        add_heading(doc, title, level=2)
        add_kv_paragraph(doc, "File", path)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    add_divider(doc)

    # ---- Section: Shared popups -------------------------------------
    add_heading(doc, "4. Shared popups & modals", level=1)

    add_heading(doc, "Sync popup (bbShowSyncPopup)", level=2)
    add_body(doc,
        "Glass card with the BrandBased B logo (breathing pulse), a label, and an animated "
        "progress bar. Used for Sync, Publish, Verify, etc. Defined in "
        "brand-modules-bundle/_shared/sync-popup.js + .css."
    )
    add_bullet(doc, "When called from inside an iframe, it postMessages the call up to the dashboard so the popup paints at the dashboard level (covers chrome + iframe together).")
    add_bullet(doc, "Phases: bar fill → done beat (\"Syncing\" → \"Synced\") → fade out.")

    add_heading(doc, "Confirm popup (bbShowConfirmPopup)", level=2)
    add_body(doc,
        "Same iframe-forward pattern as the sync popup but with a title, body, and two action "
        "buttons (Cancel + Confirm). Used by:"
    )
    add_bullet(doc, "Brands page → Delete brand. Danger variant (red inline-SVG B icon, red breathing halo, danger-red confirm button).")
    add_bullet(doc, "Preferences → Delete Account. Same danger variant.")
    add_body(doc,
        "Heavier backdrop blur (12px + saturate(1.05)) and tint (rgba(8,11,26,0.62)) than the "
        "sync popups, scoped via .bb-confirm-popup-backdrop. The popup card itself uses "
        "#ffffff78 in light mode for a soft white veil."
    )

    add_heading(doc, "Preferences modal", level=2)
    add_body(doc, "Account → Preferences opens this. Sits at z-index 200000 so it always paints above sync popups.")
    add_bullet(doc, "Top: \"Your Subscription type:\" badge — PREMIUM (blue pill) by default, FREEMIUM (amber pill) when body.bb-dev-freemium-mode is on. Both pills are in the DOM; CSS swaps which one shows.")
    add_bullet(doc, "Body: Account Email, Password, Confirm Password.")
    add_bullet(doc, "Cancel + Save buttons are equal-width, centered, 12px radius, 14px font — matching the Delete-confirm popup's action row.")
    add_bullet(doc, "Bottom: \"Delete Account?\" text link. Click opens the danger confirm popup with copy: \"If you delete your account you'll lose all current data, and any active subscription billing will be cancelled. This action can't be undone.\"")
    add_bullet(doc, "Light-mode card background: #ffffff78 (same veil as Delete confirm and Unlock Premium).")

    add_heading(doc, "Stripe Purchase Verify modal", level=2)
    add_body(doc,
        "Shares the Preferences glass styling; indigo/blue B-mark halo. Asks for email + confirm "
        "email before handing off to Stripe. Toggle-able from the dev page via "
        "localStorage[\"bbDevStripeVerifyOpen\"]."
    )

    add_heading(doc, "Brand Support modal", level=2)
    add_body(doc,
        "Opens from any [href=\"/brand-support\"] or [data-bb-brand-support] click. Reuses the "
        "Preferences glass styling; teal B-mark halo. Shows support@brandbased.ai as a clickable "
        "mailto link (blue in light mode, no white hover flash)."
    )

    add_heading(doc, "Unlock Premium gate", level=2)
    add_body(doc,
        "Full-viewport frosted glass card with the message: \"Upgrade to Premium to power your "
        "BrandBased account with AI-driven Content Commerce, intelligent automation, advanced "
        "metrics and more.\" Title \"Unlock Premium\" rendered in solid brand purple with a soft "
        "text-shadow."
    )
    add_bullet(doc, "Activated by localStorage[\"bbDevDashSimulateFreemium\"] = \"1\" plus the iframe being on a non-excluded route.")
    add_bullet(doc, "Excluded paths: any /freemium/ URL, the Start Now page, and the Brands page.")
    add_bullet(doc, "Triggered explicitly by the header Get Premium button, the Account → Subscriptions click, and the Brands → Add → Premium flow (via window.bbDevForceShowUpgradeGate() / postMessage bb-dev-show-upgrade).")
    add_bullet(doc, "Dashboard chrome (.header-nav, .nav-container, .footer-bar-brand-console) and the .account / #account-popup / .get-premium controls sit at z-index 300000+ so they paint above the gate.")

    add_divider(doc)

    # ---- Section: Freemium flow -------------------------------------
    add_heading(doc, "5. Freemium flow — end-to-end", level=1)
    add_body(doc, "Tracking the journey of a brand-new Freemium user:")
    flow = [
        ("1.", "Start Now", "Sees \"Start FREE\" CTA."),
        ("2.", "Upload Brand Assets (Free)", "Loads freemium/Freemium-Logo-upload-and-Crop-module.html. Uploads light + dark logos. The cropper modal handles both."),
        ("3.", "Brand Verification (Free)", "Loads freemium/Freemium-Meta-Verification.html. Verify button triggers the sync popup; on success the user proceeds."),
        ("4.", "Brand Settings (Free)", "Loads freemium/Freemium-Brand-Settings.html. Configure brand details. The purple side panel shows Freemium-only links."),
        ("5.", "Theme Styles (Free)", "Loads freemium/Freemium-Theme-Styles.html. Single slot per brand. Sets accent / background / typography / etc."),
        ("6.", "Hotspots (Free)", "Loads freemium/Freemium-Hotspots.html. Router force-reloads the iframe once on arrival so the page lands at the top."),
        ("7.", "Brands page", "User can manage their Freemium brand entries (up to 12). Add Brand → Freemium flow restarts at step 2."),
    ]
    for i, label, body in flow:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_styled_run(p, f"{i} ", bold=True, color=BRAND_PURPLE, size=11)
        add_styled_run(p, f"{label} — ", bold=True, size=11)
        add_styled_run(p, body, size=11)

    add_body(doc,
        "Premium-only routes (Products, AI Logic, Metrics, Payments, Scheduling, Platforms, Ad "
        "Campaigns) load behind the Unlock Premium gate. The user sees the frosted overlay with "
        "the upgrade CTA instead of the actual module.",
        italic=True, color=MUTED, size=10,
    )

    add_divider(doc)

    # ---- Section: Premium flow --------------------------------------
    add_heading(doc, "6. Premium flow — end-to-end", level=1)
    flow_p = [
        ("1.", "Start Now", "Sees \"Go Premium\" CTA (mirrors Start FREE flow)."),
        ("2.", "Upload Brand Assets", "Loads logo-upload/Logo-upload-and-Crop-module.html. Same uploader + cropper as Freemium, no /freemium/ URL constraint."),
        ("3.", "Brand Verification", "Loads logo-upload/Meta-Verification.html. Same Verify button + sync popup."),
        ("4.", "Brand Settings", "Loads Brand-Settings-Module.html. Full settings surface."),
        ("5.", "Theme Styles", "Loads Brand-Theme-Settings-Module.html. 7 slots. Active slot drives both the popup modal and Products."),
        ("6.", "Products", "Loads products/Products.html. 7 slots, mirroring Theme Styles indices."),
        ("7.", "Hotspots", "Loads hotspots/Hotspots.html. Interactive product hotspots across theme + product assets."),
        ("8.", "AI Logic / Metrics / Payments / Platforms / Scheduling / Ad Campaigns", "All accessible from the sidebar."),
        ("9.", "Brands page", "Up to 12 Premium brands (and 12 Freemium brands) per account. Premium pill visible when the dev plan-tier preview is on."),
    ]
    for i, label, body in flow_p:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_styled_run(p, f"{i} ", bold=True, color=BRAND_PURPLE, size=11)
        add_styled_run(p, f"{label} — ", bold=True, size=11)
        add_styled_run(p, body, size=11)

    add_divider(doc)

    # ---- Section: Dev tools -----------------------------------------
    add_heading(doc, "7. Dev tools — brand-modules-bundle/dev/export-settings.html", level=1)
    add_body(doc,
        "Used to (a) export a customer's persisted dashboard state into a JSON or PHP bundle for "
        "the backend, and (b) preview UI states (verification fail, Freemium dashboard, Stripe "
        "verify modal, plan-tier pills) without breaking the real flows."
    )

    add_heading(doc, "Exports", level=2)
    add_two_col_table(
        doc,
        ["Bundle key", "What it holds"],
        [
            ("brand", "Light + dark logo data URLs, last viewed mode, Asset-Lab post-crop SVG."),
            ("themes.activeSlot + themes.slots[7]", "7 theme slots; only populated slots have data. Drives both Premium and Freemium popup modals."),
            ("premium.products.slots[7]", "Premium-only. 7 product slots, indexed against the same active slot as themes."),
            ("freemium", "Note for the backend: Freemium renders themes[activeSlot] + brand. Products N/A."),
            ("other", "Catch-all for any bb* localStorage key not explicitly tracked. Backend reviews before persisting."),
        ],
    )

    add_heading(doc, "Dev toggles", level=2)
    add_two_col_table(
        doc,
        ["Toggle", "Key", "Effect"],
        [
            ("Brand Verification preview state", "bbDevVerifyState", "\"fail\" → Brand Verification pages render the red / cross / \"Unable to Verify Brand\" failure state. Removed → default Verified state."),
            ("Brand Console — dashboard tier preview", "bbDevDashSimulateFreemium", "\"1\" → Freemium simulation: Unlock Premium gate, swapped chrome, narrower module set. Removed → Premium."),
            ("Stripe purchase verify — preview modal", "bbDevStripeVerifyOpen", "Pops the Stripe verify modal on the dashboard from another tab."),
            ("Brands page — Plan tier preview", "bbDevShowPlanTier", "Shows the Premium pill on each brand circle on the Brands page."),
        ],
    )

    add_body(doc,
        "All dev-only keys are intentionally omitted from the JSON / PHP bundle so they can't "
        "leak into production data.",
        italic=True, color=MUTED, size=10,
    )

    add_divider(doc)

    # ---- Section: File map ------------------------------------------
    add_heading(doc, "8. File map — where things live", level=1)
    add_two_col_table(
        doc,
        ["Area", "Path"],
        [
            ("Dashboard shell", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-console-final/brand-console-dashboard.html"),
            ("Dashboard styles", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-console-final/css/brand-console-styles.css"),
            ("Dashboard router", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-console-final/js/dashboard-router.js"),
            ("Freemium gate", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-console-final/js/bb-dev-freemium-gate.js"),
            ("Shared sync + confirm popup", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-modules-bundle/_shared/sync-popup.{js,css}"),
            ("Shared UI", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-modules-bundle/bb-shared-ui.{css,js}"),
            ("Module pages", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-modules-bundle/<module>/"),
            ("Freemium pages", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-modules-bundle/freemium/"),
            ("Dev tools", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/brand-modules-bundle/dev/"),
            ("Standalone Premium/Freemium modals", "BB-Entire-Frontend-Modals-and-Brand-Dashbaord/premium-and-freemium-bb-modals-standalone/"),
        ],
    )

    add_divider(doc)

    # ---- Section: Theming / display modes ---------------------------
    add_heading(doc, "9. Theming — dark vs light mode", level=1)
    add_body(doc,
        "The dashboard toggles between dark (default, no class on <body>) and light "
        "(body.light-mode). Most pages also support this; the brand-pages-isolated/* copies are "
        "frozen vendor styles used by the protected Brand Settings + Brand Theme Settings pages."
    )
    add_bullet(doc, "Dark mode: .content-block uses background: #000000a1 / background-color: #000000cc.")
    add_bullet(doc, "Dark mode: shard gradient (.light-shard-ui-theme) uses linear-gradient(259.53deg, #0A3FFF 6.53%, #0056ff85 95.34%) — all blue.")
    add_bullet(doc, "Light mode: shard gradient keeps the orange→blue base (linear-gradient(259.53deg, #0A3FFF 6.53%, #F55F0A 95.34%)).")
    add_bullet(doc, "Light mode: Account chip filter: brightness(0.5); #account-title / .account p / .account h2 use #8b8b8bad.")
    add_bullet(doc, "Light mode: Logout button has a 2px solid #000 border and black text; hover swaps to a grey background with white text + outline (no blur).")
    add_bullet(doc, "Light mode: popup cards (Preferences, Delete confirm, Unlock Premium) use background: #ffffff78.")

    add_divider(doc)

    # ---- Closing ----------------------------------------------------
    add_heading(doc, "10. Glossary", level=1)
    add_two_col_table(
        doc,
        ["Term", "Meaning"],
        [
            ("Module", "A single iframe page hosted inside the dashboard's .content-block (Brands, Theme Styles, Products, …)."),
            ("Route", "An entry in dashboard-router.js's ROUTES map. Maps a route id to the module's URL + title."),
            ("Alias", "A URL pattern in ROUTE_ALIASES that isn't a full route but maps back to one (so sub-flow pages still highlight a sidebar item)."),
            ("Slot", "A persisted Theme Styles or Products configuration. Premium has 7 slots per brand; the active one drives the popup modal."),
            ("Active slot", "The slot index the customer last viewed. Stored separately for Theme Styles and Products; they're kept in lockstep by design."),
            ("Hard refresh", "Reloading the entire dashboard at ?page=<id> instead of swapping the iframe in place. Used for pages with inner-scroll layouts (Theme Styles, Products)."),
            ("Glass card / glass popup", "The frosted-translucent surface used by sync popups, confirm popups, Preferences, Stripe verify, Brand Support, Unlock Premium."),
        ],
    )

    return doc


if __name__ == "__main__":
    out = BB_ROOT / "BrandBased-Product-Handbook.docx"
    doc = build()
    doc.save(out)
    print(f"wrote: {out}  ({out.stat().st_size:,} bytes)")
